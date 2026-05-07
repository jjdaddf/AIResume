import json, os, sys
sys.path.insert(0, r'd:\projects\AIResume\backend')

from app.models.schemas import ResumeData
from app.services.export_service import ExportService
import asyncio

os.chdir(r'd:\projects\AIResume\backend')
with open('test_data.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

resume = ResumeData(**data['resume_data'])
svc = ExportService()

docx_data = svc.generate_docx(resume)
with open('exports/test_user.docx', 'wb') as f:
    f.write(docx_data)
print('DOCX saved:', len(docx_data), 'bytes')

pdf_data = asyncio.run(svc.generate_pdf(resume))
with open('exports/test_user.pdf', 'wb') as f:
    f.write(pdf_data)
print('PDF saved:', len(pdf_data), 'bytes')

# 检查 docx 结构
from docx import Document
doc = Document('exports/test_user.docx')
print('Tables:', len(doc.tables))
print('Paragraphs:', len(doc.paragraphs))
for i, s in enumerate(doc.sections):
    print(f'Section {i}: page={s.page_width/914400*25.4:.1f}x{s.page_height/914400*25.4:.1f}mm')
