"""AI 简历路由 - 润色、JD匹配、技能推荐、面试题、Word解析"""

import io
import logging
from fastapi import APIRouter, HTTPException, UploadFile, File

from app.models.schemas import (
    AIPolishRequest,
    AIPolishResponse,
    JDMatchRequest,
    JDMatchResponse,
    MatchResult,
    SkillRecommendRequest,
    SkillRecommendResponse,
    InterviewRequest,
    InterviewResponse,
    InterviewQuestion,
    ParseDocxResponse,
    PersonalInfo,
    ResumeSection,
)
from app.services.qwen_service import qwen_service
from app.services.qdrant_service import qdrant_service

router = APIRouter(prefix="/ai", tags=["AI 智能服务"])
logger = logging.getLogger(__name__)


def _build_resume_summary(resume_data) -> str:
    """将简历数据转为文本摘要供 AI 使用"""
    parts = []
    info = resume_data.personal_info
    if info.name:
        parts.append(f"姓名: {info.name}")
    if info.summary:
        parts.append(f"个人简介: {info.summary}")

    for section in resume_data.sections:
        if not section.visible:
            continue
        section_text = f"\n【{section.title}】\n"
        data = section.data
        if isinstance(data, dict):
            for k, v in data.items():
                if v and k not in ("id",):
                    section_text += f"  {k}: {v}\n"
        elif isinstance(data, str):
            section_text += f"  {data}\n"
        parts.append(section_text)

    return "\n".join(parts)


@router.post("/polish", response_model=AIPolishResponse)
async def ai_polish(request: AIPolishRequest):
    """AI 简历经历润色 - STAR 法则扩写"""
    try:
        result = await qwen_service.polish_experience(
            text=request.text,
            section_type=request.section_type,
        )
        return AIPolishResponse(
            original_text=request.text,
            polished_text=result.get("polished_text", request.text),
            star_breakdown=result.get("star_breakdown"),
            suggestions=result.get("suggestions", []),
        )
    except Exception as e:
        logger.error(f"AI 润色失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 润色服务异常: {str(e)}")


@router.post("/match-jd", response_model=JDMatchResponse)
async def match_jd(request: JDMatchRequest):
    """岗位 JD 精准匹配 (RAG)"""
    try:
        resume_summary = _build_resume_summary(request.resume_data)

        # 1. Qdrant 向量匹配
        resume_texts = []
        for section in request.resume_data.sections:
            if section.visible and section.data:
                data = section.data
                if isinstance(data, dict):
                    resume_texts.append(" ".join(str(v) for v in data.values() if v))
                elif isinstance(data, str):
                    resume_texts.append(data)

        vector_match = qdrant_service.compute_jd_resume_match(
            jd_text=request.jd_text,
            resume_texts=resume_texts,
        )

        # 2. Qwen 语义匹配
        qwen_result = await qwen_service.match_jd(
            jd_text=request.jd_text,
            resume_summary=resume_summary,
        )

        # 3. 存储 JD 向量供后续检索
        qdrant_service.store_jd_vector(
            jd_id=f"jd_{hash(request.jd_text) % (2**31)}",
            jd_text=request.jd_text,
        )

        # 合并结果
        matched_skills = [
            MatchResult(
                skill_name=s.get("skill_name", ""),
                match_type=s.get("match_type", "partial"),
                relevance_score=s.get("relevance_score", 0.0),
            )
            for s in qwen_result.get("matched_skills", [])
        ]

        return JDMatchResponse(
            overall_score=qwen_result.get("overall_score", int(vector_match["overall_similarity"] * 100)),
            matched_skills=matched_skills,
            missing_skills=qwen_result.get("missing_skills", []),
            suggestions=qwen_result.get("suggestions", []),
        )
    except Exception as e:
        logger.error(f"JD 匹配失败: {e}")
        raise HTTPException(status_code=500, detail=f"JD 匹配服务异常: {str(e)}")


@router.post("/recommend-skills", response_model=SkillRecommendResponse)
async def recommend_skills(request: SkillRecommendRequest):
    """智能技能图谱推荐"""
    try:
        # 1. Qdrant 向量检索相似技能
        vector_skills = qdrant_service.search_similar_skills(
            query_text=request.job_title,
            limit=15,
        )

        # 2. Qwen 语义推荐
        qwen_result = await qwen_service.recommend_skills(
            job_title=request.job_title,
            current_skills=request.current_skills,
        )

        # 合并去重
        existing_names = set(request.current_skills)
        recommended = qwen_result.get("recommended_skills", [])

        # 补充向量检索结果
        for vs in vector_skills:
            name = vs.get("skill_name", "")
            if name and name not in existing_names:
                recommended.append({
                    "name": name,
                    "category": vs.get("category", "其他"),
                    "relevance": vs.get("score", 0.0),
                    "description": vs.get("description", ""),
                })
                existing_names.add(name)

        return SkillRecommendResponse(
            recommended_skills=recommended[:15],
            categories=qwen_result.get("categories", {}),
        )
    except Exception as e:
        logger.error(f"技能推荐失败: {e}")
        raise HTTPException(status_code=500, detail=f"技能推荐服务异常: {str(e)}")


@router.post("/interview", response_model=InterviewResponse)
async def generate_interview(request: InterviewRequest):
    """模拟面试题生成"""
    try:
        resume_summary = _build_resume_summary(request.resume_data)
        result = await qwen_service.generate_interview_questions(
            resume_summary=resume_summary,
            jd_text=request.jd_text,
            num_questions=request.num_questions,
        )

        questions = [
            InterviewQuestion(
                question=q.get("question", ""),
                category=q.get("category", "综合"),
                answer_guide=q.get("answer_guide", ""),
                difficulty=q.get("difficulty", "medium"),
            )
            for q in result.get("questions", [])
        ]

        return InterviewResponse(questions=questions)
    except Exception as e:
        logger.error(f"面试题生成失败: {e}")
        raise HTTPException(status_code=500, detail=f"面试题生成服务异常: {str(e)}")


@router.post("/parse-docx", response_model=ParseDocxResponse)
async def parse_docx(file: UploadFile = File(...)):
    """上传 Word 文档，AI 识别文本并提取结构化简历数据"""
    # 1. 校验文件类型
    filename = file.filename or ""
    if not filename.lower().endswith((".docx", ".doc")):
        raise HTTPException(status_code=400, detail="仅支持 .docx / .doc 格式文件")

    # 2. 读取文件内容
    content = await file.read()
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="文件为空")
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小不能超过 10MB")

    # 3. 提取文本 - 支持多种方式
    raw_text = ""
    is_doc_format = filename.lower().endswith(".doc") and not filename.lower().endswith(".docx")

    # 方式1：如果是 .doc 旧格式，尝试用 subprocess 调 antiword/textract
    if is_doc_format:
        # .doc 是二进制格式，python-docx 不支持
        # 尝试用 textract 或直接读取纯文本
        try:
            import subprocess
            # 尝试 antiword
            tmp_path = f"/tmp/upload_{id(file)}.doc"
            with open(tmp_path, "wb") as f:
                f.write(content)
            result_sub = subprocess.run(["antiword", tmp_path], capture_output=True, text=True, timeout=10)
            if result_sub.returncode == 0 and result_sub.stdout.strip():
                raw_text = result_sub.stdout.strip()
            import os
            os.remove(tmp_path)
        except Exception:
            pass

        # antiword 不可用时，尝试读取纯文本（去除二进制字符）
        if not raw_text.strip():
            try:
                # 尝试 UTF-8 和 GBK 解码
                for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
                    try:
                        decoded = content.decode(encoding, errors="ignore")
                        # 过滤不可打印字符，保留中文和可打印ASCII
                        import re
                        cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", decoded)
                        # 提取有意义的行（含中文或至少3个连续ASCII字母）
                        lines = []
                        for line in cleaned.split("\n"):
                            line = line.strip()
                            if not line:
                                continue
                            # 有中文字符 或 3+连续英文字母 的行
                            if re.search(r"[\u4e00-\u9fff]", line) or re.search(r"[a-zA-Z]{3,}", line):
                                lines.append(line)
                        if lines:
                            raw_text = "\n".join(lines)
                            break
                    except Exception:
                        continue
            except Exception as e:
                logger.warning(f"从 .doc 二进制提取文本失败: {e}")

    # 方式2：.docx 格式用 python-docx 提取
    if not is_doc_format:
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))

            # 提取段落文本
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # 提取表格内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        tables_text.append(row_text)

            # 提取页眉页脚
            header_footer_text = []
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    for p in section.header.paragraphs:
                        if p.text.strip():
                            header_footer_text.append(p.text.strip())
                if section.footer and section.footer.paragraphs:
                    for p in section.footer.paragraphs:
                        if p.text.strip():
                            header_footer_text.append(p.text.strip())

            # 提取文本框内容 (shape 内的文本)
            shape_texts = []
            try:
                from docx.oxml.ns import qn
                for shape in doc.element.body.iter(qn('w:txbxContent')):
                    for p in shape.iter(qn('w:t')):
                        if p.text and p.text.strip():
                            shape_texts.append(p.text.strip())
            except Exception:
                pass

            all_parts = []
            if paragraphs:
                all_parts.append("\n".join(paragraphs))
            if tables_text:
                all_parts.append("\n--- 表格内容 ---\n" + "\n".join(tables_text))
            if header_footer_text:
                all_parts.append("\n--- 页眉页脚 ---\n" + "\n".join(header_footer_text))
            if shape_texts:
                all_parts.append("\n--- 文本框 ---\n" + "\n".join(shape_texts))

            raw_text = "\n".join(all_parts)
        except Exception as e:
            logger.error(f"Word 文档解析失败: {e}")
            # python-docx 解析失败时，尝试从 XML 中提取文本
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                with zipfile.ZipFile(io.BytesIO(content)) as z:
                    if "word/document.xml" in z.namelist():
                        with z.open("word/document.xml") as f:
                            tree = ET.parse(f)
                            root = tree.getroot()
                            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                            texts = root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                            raw_text = "\n".join(t.text for t in texts if t.text and t.text.strip())
            except Exception as e2:
                logger.error(f"XML 提取也失败: {e2}")
                raise HTTPException(status_code=400, detail=f"Word 文档解析失败: {str(e)}")

    if not raw_text.strip():
        detail_msg = "文档内容为空，无法提取信息"
        if is_doc_format:
            detail_msg = "检测到 .doc 旧格式文件，建议用 Word 另存为 .docx 格式后重新上传，可大幅提升识别准确度"
        else:
            detail_msg = "未能提取到文本内容，可能原因：1) 文档内容在图片/扫描件中 2) 文档损坏 3) 内容在特殊控件中。建议将内容复制到新的 .docx 文件后重试"
        raise HTTPException(status_code=400, detail=detail_msg)

    # 4. 调用 AI 提取结构化信息
    try:
        result = await qwen_service.parse_docx(raw_text)
    except Exception as e:
        logger.error(f"AI 解析简历失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI 解析服务异常: {str(e)}")

    # 5. 构造响应
    personal_info_data = result.get("personal_info", {})
    personal_info = PersonalInfo(
        name=personal_info_data.get("name", ""),
        phone=personal_info_data.get("phone", ""),
        email=personal_info_data.get("email", ""),
        location=personal_info_data.get("location", ""),
        website=personal_info_data.get("website", ""),
        github=personal_info_data.get("github", ""),
        summary=personal_info_data.get("summary", ""),
    )

    sections = []
    for i, sec in enumerate(result.get("sections", [])):
        sec_type = sec.get("type", "custom")
        sec_data = sec.get("data", {})
        sections.append(ResumeSection(
            id=f"{sec_type}-{i+1}",
            type=sec_type,
            title=sec.get("title", "自定义"),
            order=i,
            visible=True,
            data=sec_data,
        ))

    return ParseDocxResponse(
        personal_info=personal_info,
        sections=sections,
        raw_text=raw_text[:5000],
    )
