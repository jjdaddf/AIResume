"""文档导出服务 - Word/PDF 生成"""

import io
import os
import base64
import logging
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


from app.core.config import settings
from app.models.schemas import ResumeData, ExportFormat, ThemeStyle

logger = logging.getLogger(__name__)

# ==================== 注册 PDF 中文字体 ====================

try:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    _FONTS_REGISTERED = False

    def _register_chinese_fonts():
        global _FONTS_REGISTERED
        if _FONTS_REGISTERED:
            return
        win_fonts_dir = "C:\\Windows\\Fonts"
        if not os.path.exists(win_fonts_dir):
            return
        fonts_to_register = [
            ("Microsoft YaHei", "msyh.ttc"),
            ("SimHei", "simhei.ttf"),
            ("SimSun", "simsun.ttc"),
        ]
        for family, file in fonts_to_register:
            path = os.path.join(win_fonts_dir, file)
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont(family, path))
                    logger.info(f"PDF 字体已注册: {family}")
                except Exception as e:
                    logger.warning(f"PDF 字体注册失败 {family}: {e}")
        _FONTS_REGISTERED = True

    _register_chinese_fonts()
except ImportError:
    logger.warning("reportlab 不可用，PDF 中文可能显示异常")

# ==================== 主题配色（与前端 themes.ts 完全对应） ====================

THEME_COLORS = {
    ThemeStyle.MINIMAL: {
        "primary": RGBColor(0x1a, 0x1a, 0x1a),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x33, 0x33, 0x33),
        "heading": RGBColor(0x33, 0x33, 0x33),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xd1, 0xd5, 0xdb),
        "subtitle": RGBColor(0x37, 0x41, 0x51),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x6b, 0x72, 0x80),
        "social": RGBColor(0x9c, 0xa3, 0xaf),
    },
    ThemeStyle.GEEK: {
        "primary": RGBColor(0x0e, 0x74, 0x90),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x06, 0xb6, 0xc4),
        "heading": RGBColor(0x0e, 0x74, 0x90),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0x22, 0xd3, 0xee),
        "subtitle": RGBColor(0x0e, 0x74, 0x90),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x6b, 0x72, 0x80),
        "social": RGBColor(0x9c, 0xa3, 0xaf),
    },
    ThemeStyle.BUSINESS: {
        "primary": RGBColor(0x1e, 0x3a, 0x5f),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x1e, 0x3a, 0x5f),
        "heading": RGBColor(0x1e, 0x3a, 0x5f),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0x1e, 0x3a, 0x5f),
        "subtitle": RGBColor(0x1e, 0x3a, 0x5f),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x6b, 0x72, 0x80),
        "social": RGBColor(0x9c, 0xa3, 0xaf),
    },
    ThemeStyle.ELEGANT: {
        "primary": RGBColor(0x5c, 0x3d, 0x2e),
        "secondary": RGBColor(0x5a, 0x4a, 0x3f),
        "accent": RGBColor(0x8b, 0x6f, 0x47),
        "heading": RGBColor(0x5c, 0x3d, 0x2e),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xc4, 0xa8, 0x82),
        "subtitle": RGBColor(0x6b, 0x4c, 0x3b),
        "date": RGBColor(0x8b, 0x73, 0x55),
        "contact": RGBColor(0x8b, 0x73, 0x55),
        "social": RGBColor(0xb0, 0xa0, 0x8a),
    },
    ThemeStyle.FOREST: {
        "primary": RGBColor(0x1b, 0x43, 0x32),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x40, 0x91, 0x6c),
        "heading": RGBColor(0x2d, 0x6a, 0x4f),
        "heading_style": "left-bar",
        "name_align": "left",
        "divider": RGBColor(0x95, 0xd5, 0xb2),
        "subtitle": RGBColor(0x2d, 0x6a, 0x4f),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x6b, 0x72, 0x80),
        "social": RGBColor(0x9c, 0xa3, 0xaf),
    },
    ThemeStyle.LAVENDER: {
        "primary": RGBColor(0x5b, 0x21, 0xb6),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x8b, 0x5c, 0xf6),
        "heading": RGBColor(0x6d, 0x28, 0xd9),
        "heading_style": "dot",
        "name_align": "center",
        "divider": RGBColor(0xc4, 0xb5, 0xfd),
        "subtitle": RGBColor(0x6d, 0x28, 0xd9),
        "date": RGBColor(0x7c, 0x73, 0x94),
        "contact": RGBColor(0x7c, 0x73, 0x94),
        "social": RGBColor(0xa0, 0x99, 0xb0),
    },
    ThemeStyle.SUNSET: {
        "primary": RGBColor(0x9a, 0x34, 0x12),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0xea, 0x58, 0x0c),
        "heading": RGBColor(0xc2, 0x41, 0x0c),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xfb, 0x92, 0x3c),
        "subtitle": RGBColor(0xc2, 0x41, 0x0c),
        "date": RGBColor(0x78, 0x71, 0x6c),
        "contact": RGBColor(0x78, 0x71, 0x6c),
        "social": RGBColor(0xa8, 0xa2, 0x9e),
    },
    ThemeStyle.MIDNIGHT: {
        "primary": RGBColor(0x1e, 0x1b, 0x4b),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x43, 0x38, 0xca),
        "heading": RGBColor(0x31, 0x2e, 0x81),
        "heading_style": "left-bar",
        "name_align": "left",
        "divider": RGBColor(0x63, 0x66, 0xf1),
        "subtitle": RGBColor(0x31, 0x2e, 0x81),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x6b, 0x72, 0x80),
        "social": RGBColor(0x9c, 0xa3, 0xaf),
    },
    ThemeStyle.ROSE: {
        "primary": RGBColor(0x9d, 0x17, 0x4d),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0xec, 0x48, 0x99),
        "heading": RGBColor(0xbe, 0x18, 0x5d),
        "heading_style": "dot",
        "name_align": "center",
        "divider": RGBColor(0xf9, 0xa8, 0xd4),
        "subtitle": RGBColor(0xbe, 0x18, 0x5d),
        "date": RGBColor(0x9c, 0xa3, 0xaf),
        "contact": RGBColor(0x9c, 0xa3, 0xaf),
        "social": RGBColor(0xd1, 0xd5, 0xdb),
    },
    ThemeStyle.SLATE: {
        "primary": RGBColor(0x1e, 0x29, 0x3b),
        "secondary": RGBColor(0x47, 0x55, 0x69),
        "accent": RGBColor(0x64, 0x74, 0x8b),
        "heading": RGBColor(0x33, 0x41, 0x55),
        "heading_style": "underline",
        "name_align": "left",
        "divider": RGBColor(0xcb, 0xd5, 0xe1),
        "subtitle": RGBColor(0x33, 0x41, 0x55),
        "date": RGBColor(0x94, 0xa3, 0xb8),
        "contact": RGBColor(0x94, 0xa3, 0xb8),
        "social": RGBColor(0xcb, 0xd5, 0xe1),
    },
    ThemeStyle.TEAL: {
        "primary": RGBColor(0x13, 0x4e, 0x4a),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x14, 0xb8, 0xa6),
        "heading": RGBColor(0x0f, 0x76, 0x6e),
        "heading_style": "left-bar",
        "name_align": "left",
        "divider": RGBColor(0x5e, 0xea, 0xd4),
        "subtitle": RGBColor(0x0f, 0x76, 0x6e),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x6b, 0x72, 0x80),
        "social": RGBColor(0x9c, 0xa3, 0xaf),
    },
    ThemeStyle.WINE: {
        "primary": RGBColor(0x7f, 0x1d, 0x1d),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0xb9, 0x1c, 0x1c),
        "heading": RGBColor(0x99, 0x1b, 0x1b),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xdc, 0x26, 0x26),
        "subtitle": RGBColor(0x99, 0x1b, 0x1b),
        "date": RGBColor(0x78, 0x71, 0x6c),
        "contact": RGBColor(0x78, 0x71, 0x6c),
        "social": RGBColor(0xa8, 0xa2, 0x9e),
    },
    ThemeStyle.OCEAN: {
        "primary": RGBColor(0x0c, 0x4a, 0x6e),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x02, 0x84, 0xc7),
        "heading": RGBColor(0x03, 0x69, 0xa1),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0x38, 0xbd, 0xf8),
        "subtitle": RGBColor(0x03, 0x69, 0xa1),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x6b, 0x72, 0x80),
        "social": RGBColor(0x9c, 0xa3, 0xaf),
    },
    ThemeStyle.MOSS: {
        "primary": RGBColor(0x3f, 0x62, 0x12),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x65, 0xa3, 0x0d),
        "heading": RGBColor(0x4d, 0x7c, 0x0f),
        "heading_style": "left-bar",
        "name_align": "left",
        "divider": RGBColor(0xa3, 0xe6, 0x35),
        "subtitle": RGBColor(0x4d, 0x7c, 0x0f),
        "date": RGBColor(0x78, 0x71, 0x6c),
        "contact": RGBColor(0x78, 0x71, 0x6c),
        "social": RGBColor(0xa8, 0xa2, 0x9e),
    },
    ThemeStyle.GRAPHITE: {
        "primary": RGBColor(0x0f, 0x0f, 0x0f),
        "secondary": RGBColor(0x52, 0x52, 0x52),
        "accent": RGBColor(0x40, 0x40, 0x40),
        "heading": RGBColor(0x26, 0x26, 0x26),
        "heading_style": "underline",
        "name_align": "left",
        "divider": RGBColor(0x52, 0x52, 0x52),
        "subtitle": RGBColor(0x26, 0x26, 0x26),
        "date": RGBColor(0x73, 0x73, 0x73),
        "contact": RGBColor(0x73, 0x73, 0x73),
        "social": RGBColor(0xa3, 0xa3, 0xa3),
    },
    ThemeStyle.CORAL: {
        "primary": RGBColor(0x9f, 0x12, 0x39),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0xe1, 0x1d, 0x48),
        "heading": RGBColor(0xbe, 0x12, 0x3c),
        "heading_style": "dot",
        "name_align": "center",
        "divider": RGBColor(0xfb, 0x71, 0x85),
        "subtitle": RGBColor(0xbe, 0x12, 0x3c),
        "date": RGBColor(0x9c, 0xa3, 0xaf),
        "contact": RGBColor(0x9c, 0xa3, 0xaf),
        "social": RGBColor(0xd1, 0xd5, 0xdb),
    },
    ThemeStyle.INDIGO: {
        "primary": RGBColor(0x31, 0x2e, 0x81),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x4f, 0x46, 0xe5),
        "heading": RGBColor(0x37, 0x30, 0xa3),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0x81, 0x8c, 0xf8),
        "subtitle": RGBColor(0x37, 0x30, 0xa3),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x6b, 0x72, 0x80),
        "social": RGBColor(0x9c, 0xa3, 0xaf),
    },
    ThemeStyle.SAND: {
        "primary": RGBColor(0x78, 0x35, 0x0f),
        "secondary": RGBColor(0x57, 0x53, 0x4e),
        "accent": RGBColor(0xb4, 0x53, 0x09),
        "heading": RGBColor(0x92, 0x40, 0x0e),
        "heading_style": "underline",
        "name_align": "left",
        "divider": RGBColor(0xd9, 0x77, 0x06),
        "subtitle": RGBColor(0x92, 0x40, 0x0e),
        "date": RGBColor(0x78, 0x71, 0x6c),
        "contact": RGBColor(0x78, 0x71, 0x6c),
        "social": RGBColor(0xa8, 0xa2, 0x9e),
    },
    ThemeStyle.MIST: {
        "primary": RGBColor(0x1e, 0x40, 0xaf),
        "secondary": RGBColor(0x64, 0x74, 0x8b),
        "accent": RGBColor(0x60, 0xa5, 0xfa),
        "heading": RGBColor(0x25, 0x63, 0xeb),
        "heading_style": "left-bar",
        "name_align": "left",
        "divider": RGBColor(0x93, 0xc5, 0xfd),
        "subtitle": RGBColor(0x25, 0x63, 0xeb),
        "date": RGBColor(0x94, 0xa3, 0xb8),
        "contact": RGBColor(0x94, 0xa3, 0xb8),
        "social": RGBColor(0xcb, 0xd5, 0xe1),
    },
    ThemeStyle.CHARCOAL: {
        "primary": RGBColor(0x00, 0x00, 0x00),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x11, 0x18, 0x27),
        "heading": RGBColor(0x11, 0x18, 0x27),
        "heading_style": "underline",
        "name_align": "left",
        "divider": RGBColor(0x37, 0x41, 0x51),
        "subtitle": RGBColor(0x11, 0x18, 0x27),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x6b, 0x72, 0x80),
        "social": RGBColor(0x9c, 0xa3, 0xaf),
    },
    ThemeStyle.SAGE: {
        "primary": RGBColor(0x3f, 0x62, 0x12),
        "secondary": RGBColor(0x57, 0x53, 0x4e),
        "accent": RGBColor(0x84, 0xcc, 0x16),
        "heading": RGBColor(0x4d, 0x7c, 0x0f),
        "heading_style": "dot",
        "name_align": "center",
        "divider": RGBColor(0xbe, 0xf2, 0x64),
        "subtitle": RGBColor(0x4d, 0x7c, 0x0f),
        "date": RGBColor(0xa8, 0xa2, 0x9e),
        "contact": RGBColor(0xa8, 0xa2, 0x9e),
        "social": RGBColor(0xd6, 0xd3, 0xd1),
    },
    ThemeStyle.PLUM: {
        "primary": RGBColor(0x58, 0x1c, 0x87),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x7c, 0x3a, 0xed),
        "heading": RGBColor(0x6b, 0x21, 0xa8),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xa7, 0x8b, 0xfa),
        "subtitle": RGBColor(0x6b, 0x21, 0xa8),
        "date": RGBColor(0x7c, 0x73, 0x94),
        "contact": RGBColor(0x7c, 0x73, 0x94),
        "social": RGBColor(0xa0, 0x99, 0xb0),
    },
    ThemeStyle.COPPER: {
        "primary": RGBColor(0x7c, 0x2d, 0x12),
        "secondary": RGBColor(0x57, 0x53, 0x4e),
        "accent": RGBColor(0xb4, 0x53, 0x09),
        "heading": RGBColor(0x9a, 0x34, 0x12),
        "heading_style": "left-bar",
        "name_align": "left",
        "divider": RGBColor(0xd9, 0x77, 0x06),
        "subtitle": RGBColor(0x9a, 0x34, 0x12),
        "date": RGBColor(0x78, 0x71, 0x6c),
        "contact": RGBColor(0x78, 0x71, 0x6c),
        "social": RGBColor(0xa8, 0xa2, 0x9e),
    },
    ThemeStyle.ICE: {
        "primary": RGBColor(0x0c, 0x4a, 0x6e),
        "secondary": RGBColor(0x64, 0x74, 0x8b),
        "accent": RGBColor(0x0e, 0xa5, 0xe9),
        "heading": RGBColor(0x02, 0x84, 0xc7),
        "heading_style": "none",
        "name_align": "left",
        "divider": RGBColor(0x7d, 0xd3, 0xfc),
        "subtitle": RGBColor(0x02, 0x84, 0xc7),
        "date": RGBColor(0x94, 0xa3, 0xb8),
        "contact": RGBColor(0x94, 0xa3, 0xb8),
        "social": RGBColor(0xcb, 0xd5, 0xe1),
    },
    ThemeStyle.MAPLE: {
        "primary": RGBColor(0x9a, 0x34, 0x12),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0xdc, 0x26, 0x26),
        "heading": RGBColor(0xc2, 0x41, 0x0c),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xf8, 0x71, 0x71),
        "subtitle": RGBColor(0xc2, 0x41, 0x0c),
        "date": RGBColor(0x78, 0x71, 0x6c),
        "contact": RGBColor(0x78, 0x71, 0x6c),
        "social": RGBColor(0xa8, 0xa2, 0x9e),
    },
    # ==================== 侧边栏/双栏布局主题 ====================
    ThemeStyle.NAVY_SIDEBAR: {
        "primary": RGBColor(0x1e, 0x3a, 0x5f),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0x60, 0xa5, 0xfa),
        "heading": RGBColor(0x1e, 0x3a, 0x5f),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xe2, 0xe8, 0xf0),
        "subtitle": RGBColor(0x1e, 0x3a, 0x5f),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x94, 0xa3, 0xb8),
        "social": RGBColor(0xcb, 0xd5, 0xe1),
    },
    ThemeStyle.EMERALD_SIDEBAR: {
        "primary": RGBColor(0x06, 0x5f, 0x46),
        "secondary": RGBColor(0x37, 0x41, 0x51),
        "accent": RGBColor(0x34, 0xd3, 0x99),
        "heading": RGBColor(0x06, 0x5f, 0x46),
        "heading_style": "left-bar",
        "name_align": "center",
        "divider": RGBColor(0xd1, 0xfa, 0xe5),
        "subtitle": RGBColor(0x06, 0x5f, 0x46),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0xa7, 0xf3, 0xd0),
        "social": RGBColor(0xd1, 0xfa, 0xe5),
    },
    ThemeStyle.BURGUNDY_SIDEBAR: {
        "primary": RGBColor(0x88, 0x13, 0x37),
        "secondary": RGBColor(0x37, 0x41, 0x51),
        "accent": RGBColor(0xfb, 0x71, 0x85),
        "heading": RGBColor(0x88, 0x13, 0x37),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xfe, 0xcd, 0xd3),
        "subtitle": RGBColor(0x88, 0x13, 0x37),
        "date": RGBColor(0x78, 0x71, 0x6c),
        "contact": RGBColor(0xfd, 0xa4, 0xaf),
        "social": RGBColor(0xfe, 0xcd, 0xd3),
    },
    ThemeStyle.OBSIDIAN_SIDEBAR: {
        "primary": RGBColor(0x11, 0x18, 0x27),
        "secondary": RGBColor(0x4b, 0x55, 0x63),
        "accent": RGBColor(0xf5, 0x9e, 0x0b),
        "heading": RGBColor(0x11, 0x18, 0x27),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xe5, 0xe7, 0xeb),
        "subtitle": RGBColor(0x11, 0x18, 0x27),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x9c, 0xa3, 0xaf),
        "social": RGBColor(0xd1, 0xd5, 0xdb),
    },
    ThemeStyle.VIOLET_SIDEBAR: {
        "primary": RGBColor(0x6d, 0x28, 0xd9),
        "secondary": RGBColor(0x37, 0x41, 0x51),
        "accent": RGBColor(0xa7, 0x8b, 0xfa),
        "heading": RGBColor(0x6d, 0x28, 0xd9),
        "heading_style": "dot",
        "name_align": "center",
        "divider": RGBColor(0xe9, 0xd5, 0xff),
        "subtitle": RGBColor(0x6d, 0x28, 0xd9),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0xc4, 0xb5, 0xfd),
        "social": RGBColor(0xe9, 0xd5, 0xff),
    },
    ThemeStyle.TEAL_SIDEBAR: {
        "primary": RGBColor(0x0f, 0x76, 0x6e),
        "secondary": RGBColor(0x37, 0x41, 0x51),
        "accent": RGBColor(0x2d, 0xd4, 0xbf),
        "heading": RGBColor(0x0f, 0x76, 0x6e),
        "heading_style": "left-bar",
        "name_align": "center",
        "divider": RGBColor(0x99, 0xf6, 0xe4),
        "subtitle": RGBColor(0x0f, 0x76, 0x6e),
        "date": RGBColor(0x6b, 0x72, 0x80),
        "contact": RGBColor(0x5e, 0xea, 0xd4),
        "social": RGBColor(0x99, 0xf6, 0xe4),
    },
    ThemeStyle.AMBER_TWOCOL: {
        "primary": RGBColor(0x78, 0x35, 0x0f),
        "secondary": RGBColor(0x57, 0x53, 0x4e),
        "accent": RGBColor(0xd9, 0x77, 0x06),
        "heading": RGBColor(0x92, 0x40, 0x0e),
        "heading_style": "underline",
        "name_align": "center",
        "divider": RGBColor(0xfb, 0xbf, 0x24),
        "subtitle": RGBColor(0x92, 0x40, 0x0e),
        "date": RGBColor(0x78, 0x71, 0x6c),
        "contact": RGBColor(0x78, 0x71, 0x6c),
        "social": RGBColor(0xa8, 0xa2, 0x9e),
    },
    ThemeStyle.SLATE_TWOCOL: {
        "primary": RGBColor(0x1e, 0x29, 0x3b),
        "secondary": RGBColor(0x47, 0x55, 0x69),
        "accent": RGBColor(0x64, 0x74, 0x8b),
        "heading": RGBColor(0x33, 0x41, 0x55),
        "heading_style": "left-bar",
        "name_align": "center",
        "divider": RGBColor(0x94, 0xa3, 0xb8),
        "subtitle": RGBColor(0x33, 0x41, 0x55),
        "date": RGBColor(0x94, 0xa3, 0xb8),
        "contact": RGBColor(0x94, 0xa3, 0xb8),
        "social": RGBColor(0xcb, 0xd5, 0xe1),
    },
}


# ==================== 栏位排版配置（与前端 themes.ts sectionLayout 对应） ====================

SECTION_LAYOUTS = {
    ThemeStyle.MINIMAL: {"education": "classic", "work": "classic", "project": "classic", "skill": "inline"},
    ThemeStyle.GEEK: {"education": "classic", "work": "timeline", "project": "featured", "skill": "tags"},
    ThemeStyle.BUSINESS: {"education": "card", "work": "compact", "project": "classic", "skill": "columns"},
    ThemeStyle.ELEGANT: {"education": "card", "work": "classic", "project": "classic", "skill": "tags"},
    ThemeStyle.FOREST: {"education": "classic", "work": "timeline", "project": "featured", "skill": "columns"},
    ThemeStyle.LAVENDER: {"education": "card", "work": "compact", "project": "classic", "skill": "tags"},
    ThemeStyle.SUNSET: {"education": "classic", "work": "classic", "project": "featured", "skill": "inline"},
    ThemeStyle.MIDNIGHT: {"education": "classic", "work": "timeline", "project": "featured", "skill": "tags"},
    ThemeStyle.ROSE: {"education": "card", "work": "classic", "project": "classic", "skill": "tags"},
    ThemeStyle.SLATE: {"education": "classic", "work": "compact", "project": "featured", "skill": "columns"},
    ThemeStyle.TEAL: {"education": "classic", "work": "timeline", "project": "classic", "skill": "tags"},
    ThemeStyle.WINE: {"education": "card", "work": "classic", "project": "classic", "skill": "inline"},
    ThemeStyle.OCEAN: {"education": "classic", "work": "compact", "project": "featured", "skill": "tags"},
    ThemeStyle.MOSS: {"education": "classic", "work": "timeline", "project": "classic", "skill": "columns"},
    ThemeStyle.GRAPHITE: {"education": "classic", "work": "compact", "project": "featured", "skill": "tags"},
    ThemeStyle.CORAL: {"education": "card", "work": "classic", "project": "featured", "skill": "tags"},
    ThemeStyle.INDIGO: {"education": "classic", "work": "timeline", "project": "classic", "skill": "columns"},
    ThemeStyle.SAND: {"education": "card", "work": "compact", "project": "featured", "skill": "inline"},
    ThemeStyle.MIST: {"education": "classic", "work": "timeline", "project": "featured", "skill": "tags"},
    ThemeStyle.CHARCOAL: {"education": "classic", "work": "compact", "project": "featured", "skill": "columns"},
    ThemeStyle.SAGE: {"education": "card", "work": "classic", "project": "classic", "skill": "tags"},
    ThemeStyle.PLUM: {"education": "classic", "work": "timeline", "project": "featured", "skill": "inline"},
    ThemeStyle.COPPER: {"education": "card", "work": "compact", "project": "featured", "skill": "columns"},
    ThemeStyle.ICE: {"education": "classic", "work": "timeline", "project": "featured", "skill": "tags"},
    ThemeStyle.MAPLE: {"education": "card", "work": "classic", "project": "featured", "skill": "inline"},
    # 侧边栏/双栏布局主题
    ThemeStyle.NAVY_SIDEBAR: {"education": "card", "work": "timeline", "project": "featured", "skill": "tags"},
    ThemeStyle.EMERALD_SIDEBAR: {"education": "classic", "work": "classic", "project": "classic", "skill": "columns"},
    ThemeStyle.BURGUNDY_SIDEBAR: {"education": "card", "work": "compact", "project": "featured", "skill": "tags"},
    ThemeStyle.OBSIDIAN_SIDEBAR: {"education": "classic", "work": "timeline", "project": "featured", "skill": "columns"},
    ThemeStyle.VIOLET_SIDEBAR: {"education": "card", "work": "compact", "project": "classic", "skill": "tags"},
    ThemeStyle.TEAL_SIDEBAR: {"education": "classic", "work": "timeline", "project": "featured", "skill": "columns"},
    ThemeStyle.AMBER_TWOCOL: {"education": "card", "work": "compact", "project": "featured", "skill": "tags"},
    ThemeStyle.SLATE_TWOCOL: {"education": "classic", "work": "timeline", "project": "classic", "skill": "columns"},
}


# ==================== 字体映射（前端 themes.ts fontFamily → Word 字体） ====================

FONT_MAP_STR_TO_FAMILY = {
    "Inter": "Microsoft YaHei",
    "Helvetica Neue": "Microsoft YaHei",
    "JetBrains Mono": "Consolas",
    "Fira Code": "Consolas",
    "Georgia": "Times New Roman",
    "Noto Serif SC": "SimSun",
    "Source Han Serif": "SimSun",
    "monospace": "Consolas",
    "serif": "SimSun",
    "sans-serif": "Microsoft YaHei",
}

# 主题 ID → 字号（px，与前端的 themes.ts 一致）
# px → pt 转换: pt = px * 72 / 96 = px * 0.75
THEME_SIZES = {
    # (name_size_pt, heading_size_pt)
    ThemeStyle.MINIMAL: (18, 10),      # 24px→18pt, 13px→10pt
    ThemeStyle.GEEK: (18, 10),
    ThemeStyle.BUSINESS: (18, 10),
    ThemeStyle.ELEGANT: (17, 10),      # 22px→16.5pt→17pt
    ThemeStyle.FOREST: (18, 10),
    ThemeStyle.LAVENDER: (17, 10),
    ThemeStyle.SUNSET: (18, 10),
    ThemeStyle.MIDNIGHT: (17, 10),
    ThemeStyle.ROSE: (17, 10),
    ThemeStyle.SLATE: (17, 9),         # headingSize=12px→9pt
    ThemeStyle.TEAL: (18, 10),
    ThemeStyle.WINE: (17, 10),
    ThemeStyle.OCEAN: (18, 10),
    ThemeStyle.MOSS: (17, 10),
    ThemeStyle.GRAPHITE: (18, 10),
    ThemeStyle.CORAL: (17, 10),
    ThemeStyle.INDIGO: (17, 10),
    ThemeStyle.SAND: (17, 10),
    ThemeStyle.MIST: (17, 10),
    ThemeStyle.CHARCOAL: (18, 9),      # headingSize=12px→9pt
    ThemeStyle.SAGE: (17, 10),
    ThemeStyle.PLUM: (17, 10),
    ThemeStyle.COPPER: (17, 10),
    ThemeStyle.ICE: (17, 9),           # headingSize=12px→9pt
    ThemeStyle.MAPLE: (17, 10),
    ThemeStyle.NAVY_SIDEBAR: (17, 10),
    ThemeStyle.EMERALD_SIDEBAR: (17, 10),
    ThemeStyle.BURGUNDY_SIDEBAR: (17, 10),
    ThemeStyle.OBSIDIAN_SIDEBAR: (17, 10),
    ThemeStyle.VIOLET_SIDEBAR: (17, 10),
    ThemeStyle.TEAL_SIDEBAR: (17, 10),
    ThemeStyle.AMBER_TWOCOL: (18, 10),
    ThemeStyle.SLATE_TWOCOL: (18, 10),
}


def _get_theme_font_sizes(theme: ThemeStyle) -> tuple:
    """获取主题的字号（pt），返回 (name_pt, heading_pt)"""
    return THEME_SIZES.get(theme, (18, 10))
THEME_FONTS = {
    ThemeStyle.MINIMAL: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.GEEK: "'JetBrains Mono', 'Fira Code', monospace",
    ThemeStyle.BUSINESS: "'Georgia', 'Times New Roman', serif",
    ThemeStyle.ELEGANT: "'Noto Serif SC', 'Source Han Serif', serif",
    ThemeStyle.FOREST: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.LAVENDER: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.SUNSET: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.MIDNIGHT: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.ROSE: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.SLATE: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.TEAL: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.WINE: "'Noto Serif SC', 'Source Han Serif', serif",
    ThemeStyle.OCEAN: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.MOSS: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.GRAPHITE: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.CORAL: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.INDIGO: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.SAND: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.MIST: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.CHARCOAL: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.SAGE: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.PLUM: "'Noto Serif SC', 'Source Han Serif', serif",
    ThemeStyle.COPPER: "'Noto Serif SC', 'Source Han Serif', serif",
    ThemeStyle.ICE: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.MAPLE: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.NAVY_SIDEBAR: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.EMERALD_SIDEBAR: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.BURGUNDY_SIDEBAR: "'Noto Serif SC', 'Source Han Serif', serif",
    ThemeStyle.OBSIDIAN_SIDEBAR: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.VIOLET_SIDEBAR: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.TEAL_SIDEBAR: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.AMBER_TWOCOL: "'Inter', 'Helvetica Neue', sans-serif",
    ThemeStyle.SLATE_TWOCOL: "'Inter', 'Helvetica Neue', sans-serif",
}


def _resolve_font(theme: ThemeStyle) -> str:
    """根据主题解析为 Word 系统字体"""
    font_str = THEME_FONTS.get(theme, "'Inter', 'Helvetica Neue', sans-serif")
    for key, val in FONT_MAP_STR_TO_FAMILY.items():
        if key.lower() in font_str.lower():
            return val
    return "Microsoft YaHei"


# ==================== 工具函数 ====================

def _rgb_hex(c) -> str:
    """RGBColor 转 hex 字符串（不含 #）"""
    return str(c)


def _set_cell_border(cell, **kwargs):
    """设置表格单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val.get("sz", "0")}" '
                f'w:space="{val.get("space", "0")}" w:color="{val.get("color", "auto")}"/>'
            )
            tcBorders.append(element)
        else:
            element.set(qn("w:val"), val["val"])
            element.set(qn("w:sz"), str(val.get("sz", "0")))
            element.set(qn("w:space"), str(val.get("space", "0")))
            element.set(qn("w:color"), val.get("color", "auto"))


def _remove_table_borders(table):
    """移除表格所有边框"""
    no_border = {"val": "none", "sz": "0", "space": "0", "color": "auto"}
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}/>')
        tblPr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="none" w:sz="0" w:space="0" w:color="auto"/>')
            borders.append(el)
        else:
            el.set(qn("w:val"), "none")
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, top=no_border, left=no_border, bottom=no_border, right=no_border)


# ==================== 导出服务 ====================

class ExportService:
    """文档导出服务"""

    def _apply_layout(self, doc: Document, resume: ResumeData):
        """应用排版参数（layout 中 margin 单位是 mm，需转 cm）"""
        layout = resume.layout
        for section in doc.sections:
            section.top_margin = Cm(layout.margin_top / 10)
            section.bottom_margin = Cm(layout.margin_bottom / 10)
            section.left_margin = Cm(layout.margin_left / 10)
            section.right_margin = Cm(layout.margin_right / 10)

    def _content_width_cm(self, resume: ResumeData) -> float:
        """计算内容区域宽度（cm）"""
        layout = resume.layout
        # margin_left/right 是 mm，需转为 cm
        return 21.0 - (layout.margin_left + layout.margin_right) / 10

    # ---- card 容器（匹配前端 cardStyle） ----

    def _open_card(self, doc: Document, colors: dict, resume: ResumeData):
        """打开卡片容器（返回 cardStyle，供 _close_card 使用）"""
        card_style = colors.get("card_style", "none")
        if card_style == "none":
            return None
        cw = self._content_width_cm(resume)

        if card_style in ("bordered", "tinted"):
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(table)
            table.autofit = False
            table.width = Cm(cw)
            cell = table.cell(0, 0)
            if card_style == "bordered":
                _set_cell_border(cell, top={"val": "single", "sz": "4", "color": _rgb_hex(colors.get("heading_border_color", colors["accent"])) + "40"},
                                 bottom={"val": "single", "sz": "4", "color": _rgb_hex(colors.get("heading_border_color", colors["accent"])) + "40"},
                                 left={"val": "single", "sz": "4", "color": _rgb_hex(colors.get("heading_border_color", colors["accent"])) + "40"},
                                 right={"val": "single", "sz": "4", "color": _rgb_hex(colors.get("heading_border_color", colors["accent"])) + "40"})
            elif card_style == "tinted":
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{_rgb_hex(colors.get("tag_bg", colors.get("accent"))) + "18"}" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
            return {"table": table, "cell": cell}

        if card_style == "accent-left":
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(table)
            table.autofit = False
            table.width = Cm(cw)
            bar_cell = table.cell(0, 0)
            bar_cell.width = Cm(0.2)
            acc = _rgb_hex(colors["accent"])
            _set_cell_border(bar_cell, left={"val": "single", "sz": "18", "color": acc})
            # 浅背景
            content_cell = table.cell(0, 1)
            content_cell.width = Cm(cw - 0.2)
            tag_bg = colors.get("tag_bg", "")
            if tag_bg:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{tag_bg.replace("#", "")}" w:val="clear"/>'
                )
                content_cell._tc.get_or_add_tcPr().append(shading)
            return {"table": table, "wrapper_cell": content_cell}

        if card_style == "shadow":
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(table)
            table.autofit = False
            table.width = Cm(cw)
            cell = table.cell(0, 0)
            acc = _rgb_hex(colors.get("accent", "40916c"))
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{acc}15" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)
            _set_cell_border(cell, top={"val": "single", "sz": "2", "color": acc + "25"},
                             bottom={"val": "single", "sz": "2", "color": acc + "25"},
                             left={"val": "single", "sz": "2", "color": acc + "25"},
                             right={"val": "single", "sz": "2", "color": acc + "25"})
            return {"table": table, "cell": cell}

        return None

    def _close_card(self, card_ctx):
        """关闭卡片容器（Word 中卡片容器在 _open_card 已创建完成）"""
        pass

    def _add_item_separator(self, doc: Document, colors: dict, resume: ResumeData, in_card=False):
        """在同类条目之间添加分隔线（匹配前端 mt-3 pt-3 ≈ 24px ≈ Pt(18) 间距）"""
        div = _rgb_hex(colors.get("divider", colors["accent"]))
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 0.3  # 最小行高
        run = para.add_run("")
        run.font.size = Pt(2)
        pPr = para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="2" w:color="{div}40"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    # ---- Word: 个人信息 ----    
    def _add_personal_info(self, doc: Document, resume: ResumeData, colors: dict, theme_font: str = "Microsoft YaHei", name_pt: int = 18):
        """添加个人信息区域（匹配前端 headerStyle + nameTransform）"""
        info = resume.personal_info
        if not info.name:
            return

        name_transform = colors.get("name_transform", "none")
        name_align_center = colors.get("name_align", "center") == "center"
        align = WD_ALIGN_PARAGRAPH.CENTER if name_align_center else WD_ALIGN_PARAGRAPH.LEFT

        def _apply_name_transform(text: str) -> str:
            if name_transform == "uppercase":
                return text.upper()
            elif name_transform == "wide-spacing":
                return " ".join(list(text))
            return text

        display_name = _apply_name_transform(info.name)

        if info.photo:
            # 照片 + 信息 两列表格
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _remove_table_borders(table)
            table.autofit = False

            content_w = self._content_width_cm(resume)
            table.width = Cm(content_w)
            left_cell = table.cell(0, 0)
            right_cell = table.cell(0, 1)
            left_cell.width = Cm(3.0)
            right_cell.width = Cm(content_w - 3.0)

            # 照片
            try:
                photo_data = info.photo
                if photo_data.startswith("data:"):
                    _, data = photo_data.split(",", 1)
                    image_bytes = base64.b64decode(data)
                else:
                    image_bytes = base64.b64decode(photo_data)
                photo_para = left_cell.paragraphs[0]
                photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                photo_run = photo_para.add_run()
                photo_run.add_picture(io.BytesIO(image_bytes), width=Cm(2.5), height=Cm(3.5))
            except Exception as e:
                logger.warning(f"照片处理失败: {e}")

            # 右侧信息
            right_cell.paragraphs[0].text = ""
            # 姓名
            name_para = right_cell.paragraphs[0]
            name_run = name_para.add_run(display_name)
            name_run.bold = True
            name_run.font.size = Pt(name_pt)
            name_run.font.color.rgb = colors["primary"]
            name_run.font.name = theme_font
            name_run.element.rPr.rFonts.set(qn("w:eastAsia"), theme_font)

            # 联系方式（前端 text-sm = 14px → Pt(11)）
            contact_parts = [p for p in [info.phone, info.email, info.location] if p]
            if contact_parts:
                cp = right_cell.add_paragraph()
                cr = cp.add_run(" | ".join(contact_parts))
                cr.font.size = Pt(11)
                cr.font.color.rgb = colors.get("contact", colors["secondary"])

            # 社交（前端 text-xs = 12px → Pt(9)）
            social_parts = []
            if info.website: social_parts.append(f"网站: {info.website}")
            if info.linkedin: social_parts.append(f"LinkedIn: {info.linkedin}")
            if info.github: social_parts.append(f"GitHub: {info.github}")
            if social_parts:
                sp = right_cell.add_paragraph()
                sr = sp.add_run(" | ".join(social_parts))
                sr.font.size = Pt(9)
                sr.font.color.rgb = colors.get("social", colors["secondary"])

            # 简介（前端 text-sm = 14px → Pt(11)）
            if info.summary:
                sump = right_cell.add_paragraph()
                sumr = sump.add_run(info.summary)
                sumr.font.size = Pt(11)
                sumr.font.color.rgb = colors["secondary"]
        else:
            name_para = doc.add_paragraph()
            name_para.alignment = align
            name_run = name_para.add_run(display_name)
            name_run.bold = True
            name_run.font.size = Pt(name_pt)
            name_run.font.color.rgb = colors["primary"]
            name_run.font.name = theme_font
            name_run.element.rPr.rFonts.set(qn("w:eastAsia"), theme_font)

            contact_parts = [p for p in [info.phone, info.email, info.location] if p]
            if contact_parts:
                cp = doc.add_paragraph()
                cp.alignment = align
                cr = cp.add_run(" | ".join(contact_parts))
                cr.font.size = Pt(11)
                cr.font.color.rgb = colors.get("contact", colors["secondary"])

            social_parts = []
            if info.website: social_parts.append(f"网站: {info.website}")
            if info.linkedin: social_parts.append(f"LinkedIn: {info.linkedin}")
            if info.github: social_parts.append(f"GitHub: {info.github}")
            if social_parts:
                sp = doc.add_paragraph()
                sp.alignment = align
                sr = sp.add_run(" | ".join(social_parts))
                sr.font.size = Pt(9)
                sr.font.color.rgb = colors.get("social", colors["secondary"])

            if info.summary:
                sump = doc.add_paragraph()
                sump.alignment = align
                sumr = sump.add_run(info.summary)
                sumr.font.size = Pt(11)
                sumr.font.color.rgb = colors["secondary"]

        # 分隔线
        hr_para = doc.add_paragraph()
        hr_para.paragraph_format.space_before = Pt(4)
        hr_para.paragraph_format.space_after = Pt(4)
        pPr = hr_para._p.get_or_add_pPr()
        divider_hex = _rgb_hex(colors.get("divider", colors["accent"]))
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{divider_hex}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    # ---- Word: 栏位标题 ----
    def _add_section_title(self, doc: Document, title: str, colors: dict, resume: ResumeData = None,
                           section_type: str = "", heading_pt: int = 10):
        """添加栏位标题（匹配前端 headingStyle 和 SectionHeading 组件）"""
        style = colors.get("heading_style", "underline")
        # space_before: 组之间 mb-4=16px≈Pt(12)，去掉标题的 Pt(8) 还剩余～Pt(4)组底边距
        # space_after: 匹配前端 SectionHeading 的 mb-2=8px≈Pt(6)
        sb = Pt(8)
        sa = Pt(6)

        if style == "left-bar":
            # 左侧竖条（小方块） + 标题
            para = doc.add_paragraph()
            para.paragraph_format.space_before = sb
            para.paragraph_format.space_after = sa
            bar_run = para.add_run("■ ")
            bar_run.font.size = Pt(8)
            bar_run.font.color.rgb = colors["accent"]
            title_run = para.add_run(title.upper())
            title_run.bold = True
            title_run.font.size = Pt(heading_pt)
            title_run.font.color.rgb = colors["heading"]

        elif style == "dot":
            para = doc.add_paragraph()
            para.paragraph_format.space_before = sb
            para.paragraph_format.space_after = sa
            dot_run = para.add_run("● ")
            dot_run.font.size = Pt(8)
            dot_run.font.color.rgb = colors["accent"]
            title_run = para.add_run(title.upper())
            title_run.bold = True
            title_run.font.size = Pt(heading_pt)
            title_run.font.color.rgb = colors["heading"]

        elif style == "none":
            para = doc.add_paragraph()
            para.paragraph_format.space_before = sb
            para.paragraph_format.space_after = sa
            run = para.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(heading_pt)
            run.font.color.rgb = colors["heading"]

        else:  # underline（默认）
            para = doc.add_paragraph()
            para.paragraph_format.space_before = sb
            para.paragraph_format.space_after = sa
            run = para.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(heading_pt)
            run.font.color.rgb = colors["heading"]
            pPr = para._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="{_rgb_hex(colors["accent"])}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

    # ---- Word: 内容区段 ----
    def _add_item_header(self, doc: Document, title: str, date_str: str, colors: dict, resume: ResumeData):
        """添加条目标题行：标题(粗) + 日期(右对齐) 同一行"""
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _remove_table_borders(table)
        table.autofit = False

        content_w = self._content_width_cm(resume)
        title_cell = table.cell(0, 0)
        date_cell = table.cell(0, 1)
        table.width = Cm(content_w)
        title_cell.width = Cm(content_w * 0.65)
        date_cell.width = Cm(content_w * 0.35)

        # 标题
        tp = title_cell.paragraphs[0]
        tp.paragraph_format.space_before = Pt(2)
        tp.paragraph_format.space_after = Pt(0)
        tr = tp.add_run(title)
        tr.bold = True
        tr.font.size = Pt(11)
        tr.font.color.rgb = colors["primary"]

        # 日期
        dp = date_cell.paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dp.paragraph_format.space_before = Pt(2)
        dp.paragraph_format.space_after = Pt(0)
        dr = dp.add_run(date_str)
        dr.font.size = Pt(9)
        dr.font.color.rgb = colors.get("date", colors["secondary"])

    def _add_education(self, doc: Document, section, colors: dict, resume: ResumeData, section_layout: dict = None):
        """添加教育背景"""
        from app.models.schemas import EducationItem
        data = section.data
        if isinstance(data, dict):
            data = EducationItem(**data)
        if not isinstance(data, EducationItem):
            return

        sl = section_layout or {}
        edu_layout = sl.get("education", "classic")
        date_str = f"{data.start_date} - {data.end_date or '至今'}" if data.start_date else ""

        if edu_layout == "card":
            # 卡片式：学校+日期同行，学位专业GPA在下方，左边框
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.width = Cm(self._content_width_cm(resume))
            # 左边框列
            cell0 = table.cell(0, 0)
            cell0.width = Cm(0.2)
            _set_cell_border(cell0, left={"val": "single", "sz": "12", "color": _rgb_hex(colors["accent"])})
            _remove_table_borders(table)
            _set_cell_border(table.cell(0, 1), **{f"top": {"val": "none", "sz": "0", "color": "auto"}, f"bottom": {"val": "none", "sz": "0", "color": "auto"}, f"left": {"val": "none", "sz": "0", "color": "auto"}, f"right": {"val": "none", "sz": "0", "color": "auto"}})

            self._add_item_header(doc, data.school, date_str, colors, resume)
            detail_parts = [p for p in [data.degree, data.major] if p]
            if detail_parts:
                dp = doc.add_paragraph()
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.space_after = Pt(0)
                dr = dp.add_run(" ".join(detail_parts))
                dr.font.size = Pt(11)
                dr.font.color.rgb = colors["secondary"]
            if data.gpa:
                gp = doc.add_paragraph()
                gp.paragraph_format.space_before = Pt(0)
                gp.paragraph_format.space_after = Pt(0)
                gr = gp.add_run(f"★ GPA: {data.gpa}")
                gr.font.size = Pt(11)
                gr.font.color.rgb = colors["heading"]
        else:
            # classic
            self._add_item_header(doc, data.school, date_str, colors, resume)
            detail_parts = [p for p in [data.degree, data.major] if p]
            if detail_parts:
                dp = doc.add_paragraph()
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.space_after = Pt(0)
                dr = dp.add_run(" ".join(detail_parts))
                dr.font.size = Pt(11)
                dr.font.color.rgb = colors["secondary"]
            if data.gpa:
                gp = doc.add_paragraph()
                gp.paragraph_format.space_before = Pt(0)
                gp.paragraph_format.space_after = Pt(0)
                gr = gp.add_run(f"GPA: {data.gpa}")
                gr.font.size = Pt(11)
                gr.font.color.rgb = colors.get("date", colors["secondary"])

        for h in data.highlights:
            hp = doc.add_paragraph(style="List Bullet")
            hp.paragraph_format.space_before = Pt(0)
            hr = hp.add_run(h)
            hr.font.size = Pt(11)
            hr.font.color.rgb = colors["secondary"]

    def _add_work_experience(self, doc: Document, section, colors: dict, resume: ResumeData, section_layout: dict = None):
        """添加工作经历"""
        from app.models.schemas import WorkExperienceItem
        data = section.data
        if isinstance(data, dict):
            data = WorkExperienceItem(**data)
        if not isinstance(data, WorkExperienceItem):
            return

        sl = section_layout or {}
        work_layout = sl.get("work", "classic")
        date_str = f"{data.start_date} - {data.end_date or '至今'}" if data.start_date else ""

        if work_layout == "compact":
            # 紧凑：公司·职位 同行
            self._add_item_header(doc, f"{data.company} · {data.position}" if data.position else data.company, date_str, colors, resume)
        elif work_layout == "timeline":
            # 时间线：加左缩进
            self._add_item_header(doc, data.company, date_str, colors, resume)
            if data.position:
                pp = doc.add_paragraph()
                pp.paragraph_format.space_before = Pt(0)
                pp.paragraph_format.space_after = Pt(0)
                pp.paragraph_format.left_indent = Cm(0.5)
                pr = pp.add_run(data.position)
                pr.font.size = Pt(11)
                pr.font.color.rgb = colors.get("subtitle", colors["accent"])
        else:
            self._add_item_header(doc, data.company, date_str, colors, resume)
            if data.position:
                pp = doc.add_paragraph()
                pp.paragraph_format.space_before = Pt(0)
                pp.paragraph_format.space_after = Pt(0)
                pr = pp.add_run(data.position)
                pr.font.size = Pt(11)
                pr.font.color.rgb = colors.get("subtitle", colors["accent"])

        if data.description:
            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(1)
            dpr = dp.add_run(data.description)
            dpr.font.size = Pt(11)
            dpr.font.color.rgb = colors["secondary"]

        for a in data.achievements:
            ap = doc.add_paragraph(style="List Bullet")
            ap.paragraph_format.space_before = Pt(0)
            ar = ap.add_run(a)
            ar.font.size = Pt(11)
            ar.font.color.rgb = colors["secondary"]

    def _add_project(self, doc: Document, section, colors: dict, resume: ResumeData, section_layout: dict = None):
        """添加项目经验"""
        from app.models.schemas import ProjectItem
        data = section.data
        if isinstance(data, dict):
            data = ProjectItem(**data)
        if not isinstance(data, ProjectItem):
            return

        sl = section_layout or {}
        proj_layout = sl.get("project", "classic")
        date_str = f"{data.start_date} - {data.end_date or '至今'}" if data.start_date else ""

        self._add_item_header(doc, data.name, date_str, colors, resume)

        if data.role:
            rp = doc.add_paragraph()
            rp.paragraph_format.space_before = Pt(0)
            rp.paragraph_format.space_after = Pt(0)
            rr = rp.add_run(data.role)
            rr.font.size = Pt(11)
            rr.font.color.rgb = colors.get("subtitle", colors["accent"])

        if data.tech_stack:
            tp = doc.add_paragraph()
            tp.paragraph_format.space_before = Pt(1)
            if proj_layout == "featured":
                tr = tp.add_run("▸ 技术栈: " + " · ".join(data.tech_stack))
                tr.font.size = Pt(11)
                tr.font.color.rgb = colors["heading"]
                tr.bold = True
            else:
                tr = tp.add_run("技术栈: " + " · ".join(data.tech_stack))
                tr.font.size = Pt(9)
                tr.font.color.rgb = colors.get("subtitle", colors["accent"])

        if data.description:
            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(1)
            dpr = dp.add_run(data.description)
            dpr.font.size = Pt(11)
            dpr.font.color.rgb = colors["secondary"]

        for a in data.achievements:
            ap = doc.add_paragraph(style="List Bullet")
            ap.paragraph_format.space_before = Pt(0)
            ar = ap.add_run(a)
            ar.font.size = Pt(11)
            ar.font.color.rgb = colors["secondary"]

    def _add_skills(self, doc: Document, section, colors: dict, resume: ResumeData, section_layout: dict = None):
        """添加技能清单"""
        from app.models.schemas import SkillItem
        data = section.data
        if isinstance(data, dict):
            data = SkillItem(**data)
        if not isinstance(data, SkillItem):
            return

        sl = section_layout or {}
        skill_layout = sl.get("skill", "inline")

        if skill_layout == "columns" and len(data.skills) > 2:
            # 多列：用表格
            cols = 3
            if data.category:
                cp = doc.add_paragraph()
                cp.paragraph_format.space_before = Pt(1)
                cr = cp.add_run(data.category)
                cr.bold = True
                cr.font.size = Pt(11)
                cr.font.color.rgb = colors.get("subtitle", colors["primary"])
            rows = (len(data.skills) + cols - 1) // cols
            table = doc.add_table(rows=rows, cols=cols)
            _remove_table_borders(table)
            for i, skill in enumerate(data.skills):
                r, c = divmod(i, cols)
                cell = table.cell(r, c)
                cell.text = ""
                run = cell.paragraphs[0].add_run(f"› {skill}")
                run.font.size = Pt(11)
                run.font.color.rgb = colors["secondary"]
        elif skill_layout == "tags":
            # 标签云：用 · 分隔带括号
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(1)
            if data.category:
                cr = para.add_run(f"{data.category} ")
                cr.bold = True
                cr.font.size = Pt(11)
                cr.font.color.rgb = colors.get("subtitle", colors["primary"])
            sr = para.add_run("  ".join(f"[{s}]" for s in data.skills))
            sr.font.size = Pt(11)
            sr.font.color.rgb = colors["secondary"]
        else:
            # inline
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(1)
            if data.category:
                cr = para.add_run(f"{data.category}: ")
                cr.bold = True
                cr.font.size = Pt(11)
                cr.font.color.rgb = colors.get("subtitle", colors["primary"])
            sr = para.add_run(" · ".join(data.skills))
            sr.font.size = Pt(11)
            sr.font.color.rgb = colors["secondary"]

    # ---- Word 生成 ----
    def generate_docx(self, resume: ResumeData) -> bytes:
        """生成 Word 文档（匹配前端 ResumePreviewContent 渲染逻辑）"""
        doc = Document()
        self._apply_layout(doc, resume)
        colors = THEME_COLORS.get(resume.theme, THEME_COLORS[ThemeStyle.MINIMAL])
        theme_font = _resolve_font(resume.theme)
        name_pt, heading_pt = _get_theme_font_sizes(resume.theme)

        # 默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = theme_font
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), theme_font)
        # 清除默认段落间距（匹配前端，避免额外累积）
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        # 应用行高（匹配前端 lineHeight: store.layout.line_height）
        from docx.enum.text import WD_LINE_SPACING
        pf.line_spacing = resume.layout.line_height
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        # 列表样式也清除间距
        try:
            list_style = doc.styles["List Bullet"]
            list_style.paragraph_format.space_before = Pt(0)
            list_style.paragraph_format.space_after = Pt(0)
            list_style.paragraph_format.line_spacing = resume.layout.line_height
            list_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            lf = list_style.font
            lf.name = theme_font
            lf.size = Pt(11)
            list_style.element.rPr.rFonts.set(qn("w:eastAsia"), theme_font)
        except Exception:
            pass

        # 个人信息（传递主题字号）
        self._add_personal_info(doc, resume, colors, theme_font, name_pt)

        # 栏位：按 type 分组（匹配前端 groupedSections 逻辑）
        groups = []
        seen_types = set()
        for s in sorted(resume.sections, key=lambda x: x.order):
            if not s.visible:
                continue
            if s.type not in seen_types:
                seen_types.add(s.type)
                groups.append({"type": s.type, "title": s.title, "sections": [s]})
            else:
                for g in groups:
                    if g["type"] == s.type:
                        g["sections"].append(s)
                        break

        section_layout = SECTION_LAYOUTS.get(resume.theme, SECTION_LAYOUTS[ThemeStyle.MINIMAL])
        section_handlers = {
            "education": self._add_education,
            "work": self._add_work_experience,
            "project": self._add_project,
            "skill": self._add_skills,
        }

        for group in groups:
            # 打开卡片容器（对应前端 cardStyle）
            card_ctx = self._open_card(doc, colors, resume)
            target_doc = doc
            if card_ctx:
                wrapper_cell = card_ctx.get("wrapper_cell") or card_ctx.get("cell")
                if wrapper_cell:
                    # 写入卡片单元格内
                    inner_doc = Document()
                    inner_doc.sections[0].top_margin = Cm(0)
                    inner_doc.sections[0].bottom_margin = Cm(0)
                    inner_doc.sections[0].left_margin = Cm(0.3)
                    inner_doc.sections[0].right_margin = Cm(0.3)
                    target_doc = inner_doc
                    inner_font = inner_doc.styles["Normal"].font
                    inner_font.name = theme_font
                    inner_font.size = Pt(11)
                    inner_doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), theme_font)

            # 栏位标题（每组只显示一次，传递主题 heading 字号）
            self._add_section_title(target_doc, group["title"], colors, resume, group["type"], heading_pt)

            # 组内各条目
            for idx, section in enumerate(group["sections"]):
                if idx > 0:
                    # 同类型之间的分隔线（匹配前端 idx > 0 ? 'mt-3 pt-3 border-t'）
                    self._add_item_separator(target_doc, colors, resume)
                handler = section_handlers.get(section.type)
                if handler:
                    handler(target_doc, section, colors, resume, section_layout)
                elif section.data:
                    para = target_doc.add_paragraph()
                    text = section.data if isinstance(section.data, str) else str(section.data)
                    run = para.add_run(text)
                    run.font.size = Pt(11)
                    run.font.color.rgb = colors["secondary"]

            # 组间间距（匹配前端 mb-4=16px≈Pt(12)，但标题已有 Pt(8) space_before，这里只需补 Pt(4)）
            if group != groups[-1]:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(4)
                spacer_run = spacer.add_run("")
                spacer_run.font.size = Pt(1)

            # 关闭卡片：将 inner_doc 内容复制到卡片单元格
            if card_ctx and card_ctx.get("wrapper_cell") or card_ctx and card_ctx.get("cell"):
                wrapper_cell = card_ctx.get("wrapper_cell") or card_ctx.get("cell")
                if wrapper_cell:
                    # 清空默认段落
                    wrapper_cell.paragraphs[0].text = ""
                    # 如果 target_doc 是 inner_doc，将 inner_doc 内容合并到父文档
                    if hasattr(target_doc, 'paragraphs') and target_doc is not doc:
                        for p in target_doc.paragraphs:
                            new_p = wrapper_cell.add_paragraph(p.text if p.text else "")
                            new_p.paragraph_format.space_before = p.paragraph_format.space_before
                            new_p.paragraph_format.space_after = p.paragraph_format.space_after
                            for run in p.runs:
                                new_run = new_p.add_run(run.text if run.text else "")
                                new_run.bold = run.bold
                                new_run.font.size = run.font.size
                                new_run.font.color.rgb = run.font.color.rgb
                        # 合并 inner_doc 中的表格
                        for t in target_doc.element.body:
                            from lxml import etree
                            tag = t.tag.split('}')[-1] if '}' in t.tag else t.tag
                            if tag == 'tbl':
                                import copy
                                wrapper_cell._tc.append(copy.deepcopy(t))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ==================== PDF 生成 ====================

    def _build_html(self, resume: ResumeData) -> str:
        """生成 PDF 用的 HTML（table 布局，兼容 xhtml2pdf）"""
        from app.models.schemas import EducationItem, WorkExperienceItem, ProjectItem, SkillItem

        colors = THEME_COLORS.get(resume.theme, THEME_COLORS[ThemeStyle.MINIMAL])
        primary = f"#{_rgb_hex(colors['primary'])}"
        secondary = f"#{_rgb_hex(colors['secondary'])}"
        accent = f"#{_rgb_hex(colors['accent'])}"
        heading = f"#{_rgb_hex(colors['heading'])}"
        divider_c = f"#{_rgb_hex(colors.get('divider', colors['accent']))}"
        subtitle_c = f"#{_rgb_hex(colors.get('subtitle', colors['accent']))}"
        date_c = f"#{_rgb_hex(colors.get('date', colors['secondary']))}"
        contact_c = f"#{_rgb_hex(colors.get('contact', colors['secondary']))}"
        social_c = f"#{_rgb_hex(colors.get('social', colors['secondary']))}"

        heading_style = colors.get("heading_style", "underline")
        name_align = colors.get("name_align", "center")
        layout = resume.layout
        info = resume.personal_info
        section_layout = SECTION_LAYOUTS.get(resume.theme, SECTION_LAYOUTS[ThemeStyle.MINIMAL])

        # 栏位标题装饰
        if heading_style == "left-bar":
            heading_deco = f'<td style="width:3px;background-color:{accent};padding:0;"></td>'
            heading_td_style = 'padding-left:8px;'
        elif heading_style == "dot":
            heading_deco = f'<td style="width:14px;color:{accent};font-size:10pt;vertical-align:middle;">●</td>'
            heading_td_style = 'padding-left:4px;'
        else:
            heading_deco = ''
            heading_td_style = ''

        heading_border = f'border-bottom:2px solid {accent};' if heading_style == "underline" else ''

        # 获取 PDF 字体路径
        _font_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "exports")

        html_parts = [f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8"/>
<style>
@page {{
    size: A4;
    margin: {layout.margin_top}mm {layout.margin_right}mm {layout.margin_bottom}mm {layout.margin_left}mm;
}}
@font-face {{
    font-family: "PDFCN";
    src: url("{_font_dir}/msyh.ttc");
}}
body {{
    font-family: "PDFCN", "Microsoft YaHei", "SimHei", sans-serif;
    font-size: 10pt;
    color: {secondary};
    line-height: {layout.line_height};
}}
.photo {{ width: 25mm; height: 35mm; }}
hr {{ border: none; border-top: 1px solid {divider_c}; margin: 8px 0; }}
ul {{ margin: 2px 0; padding-left: 20px; }}
li {{ font-size: 10pt; color: {secondary}; margin-bottom: 1px; }}
table.layout {{ border-collapse: collapse; width: 100%; }}
table.layout td {{ padding: 0; vertical-align: top; border: none; }}
</style>
</head>
<body>"""]

        # 个人信息
        if info.name:
            if info.photo:
                photo_src = info.photo if info.photo.startswith("data:") else f"data:image/png;base64,{info.photo}"
                html_parts.append(f"""
<table class="layout" style="margin-bottom:8px;">
<tr>
<td style="width:30mm;padding-right:12px;"><img class="photo" src="{photo_src}"/></td>
<td style="vertical-align:top;">
<div style="font-size:22pt;font-weight:bold;color:{primary};">{info.name}</div>""")
            else:
                ta = "center" if name_align == "center" else "left"
                html_parts.append(f"""
<div style="text-align:{ta};margin-bottom:8px;">
<div style="font-size:22pt;font-weight:bold;color:{primary};">{info.name}</div>""")

            contact_items = [p for p in [info.phone, info.email, info.location] if p]
            if contact_items:
                html_parts.append(f'<div style="font-size:10pt;color:{contact_c};margin-top:4px;">{" | ".join(contact_items)}</div>')

            social_items = []
            if info.website: social_items.append(f"网站: {info.website}")
            if info.linkedin: social_items.append(f"LinkedIn: {info.linkedin}")
            if info.github: social_items.append(f"GitHub: {info.github}")
            if social_items:
                html_parts.append(f'<div style="font-size:9pt;color:{social_c};margin-top:2px;">{" | ".join(social_items)}</div>')

            if info.summary:
                html_parts.append(f'<div style="font-size:10pt;color:{secondary};margin-top:8px;">{info.summary}</div>')

            if info.photo:
                html_parts.append("</td></tr></table>")
            else:
                html_parts.append("</div>")

            html_parts.append("<hr/>")

        # 栏位
        sorted_sections = sorted(resume.sections, key=lambda s: s.order)
        for section in sorted_sections:
            if not section.visible:
                continue

            # 栏位标题
            if heading_style == "left-bar":
                html_parts.append(f"""
<table class="layout" style="margin-top:12px;margin-bottom:4px;">
<tr>{heading_deco}
<td style="{heading_td_style}"><div style="font-size:13pt;font-weight:bold;color:{heading};{heading_border}padding-bottom:2px;">{section.title}</div></td>
</tr></table>""")
            elif heading_style == "dot":
                html_parts.append(f"""
<table class="layout" style="margin-top:12px;margin-bottom:4px;">
<tr>{heading_deco}
<td style="{heading_td_style}"><div style="font-size:13pt;font-weight:bold;color:{heading};{heading_border}padding-bottom:2px;">{section.title}</div></td>
</tr></table>""")
            elif heading_style == "none":
                html_parts.append(f'<div style="font-size:13pt;font-weight:bold;color:{heading};margin-top:12px;margin-bottom:4px;">{section.title}</div>')
            else:  # underline
                html_parts.append(f'<div style="font-size:13pt;font-weight:bold;color:{heading};border-bottom:2px solid {accent};padding-bottom:2px;margin-top:12px;margin-bottom:4px;">{section.title}</div>')

            data = section.data
            if isinstance(data, dict):
                pass
            else:
                data = {"text": str(data)} if data else {}

            if section.type == "education":
                try:
                    d = EducationItem(**data) if isinstance(data, dict) else data
                    date_str = f"{d.start_date} - {d.end_date or '至今'}" if d.start_date else ""
                    edu_layout = section_layout.get("education", "classic")
                    if edu_layout == "card":
                        html_parts.append(f"""
<table class="layout" style="margin-top:4px;">
<tr><td style="width:3px;background-color:{accent};padding:0;"></td>
<td style="padding-left:8px;background-color:{accent}15;">
<table class="layout"><tr>
<td style="width:65%;"><span style="font-size:11pt;font-weight:bold;color:{primary};">{d.school}</span></td>
<td style="width:35%;text-align:right;"><span style="font-size:9pt;color:{date_c};">{date_str}</span></td>
</tr></table>""")
                        detail = " ".join([p for p in [d.degree, d.major] if p])
                        gpa_str = f' <span style="background-color:{accent}20;color:{heading};padding:1px 4px;border-radius:2px;font-size:8pt;">GPA {d.gpa}</span>' if d.gpa else ""
                        if detail or gpa_str:
                            html_parts.append(f'<div style="font-size:10pt;color:{secondary};">{detail}{gpa_str}</div>')
                        if d.highlights:
                            html_parts.append("<ul>" + "".join(f"<li>{h}</li>" for h in d.highlights) + "</ul>")
                        html_parts.append("</td></tr></table>")
                    else:
                        html_parts.append(f"""
<table class="layout" style="margin-top:4px;">
<tr>
<td style="width:65%;"><span style="font-size:11pt;font-weight:bold;color:{primary};">{d.school}</span></td>
<td style="width:35%;text-align:right;"><span style="font-size:9pt;color:{date_c};">{date_str}</span></td>
</tr></table>""")
                        detail = " ".join([p for p in [d.degree, d.major] if p])
                        gpa_str = f' | <span style="color:{date_c};">GPA: {d.gpa}</span>' if d.gpa else ""
                        if detail or gpa_str:
                            html_parts.append(f'<div style="font-size:10pt;color:{secondary};">{detail}{gpa_str}</div>')
                        if d.highlights:
                            html_parts.append("<ul>" + "".join(f"<li>{h}</li>" for h in d.highlights) + "</ul>")
                except Exception as e:
                    logger.warning(f"教育背景渲染失败: {e}")

            elif section.type == "work":
                try:
                    d = WorkExperienceItem(**data) if isinstance(data, dict) else data
                    date_str = f"{d.start_date} - {d.end_date or '至今'}" if d.start_date else ""
                    work_layout = section_layout.get("work", "classic")
                    if work_layout == "timeline":
                        html_parts.append(f"""
<table class="layout" style="margin-top:4px;">
<tr>
<td style="width:3px;vertical-align:top;padding-top:3px;">
<div style="width:3px;height:3px;background-color:{accent};border-radius:50%;"></div>
</td>
<td style="padding-left:8px;">
<table class="layout"><tr>
<td style="width:65%;"><span style="font-size:11pt;font-weight:bold;color:{primary};">{d.company}</span></td>
<td style="width:35%;text-align:right;"><span style="font-size:9pt;color:{date_c};">{date_str}</span></td>
</tr></table>""")
                        if d.position:
                            html_parts.append(f'<div style="font-size:10pt;color:{subtitle_c};">{d.position}</div>')
                        if d.description:
                            html_parts.append(f'<div style="font-size:10pt;color:{secondary};margin-top:2px;">{d.description}</div>')
                        if d.achievements:
                            html_parts.append("<ul>" + "".join(f"<li>{a}</li>" for a in d.achievements) + "</ul>")
                        html_parts.append("</td></tr></table>")
                    elif work_layout == "compact":
                        pos_part = f' · {d.position}' if d.position else ''
                        html_parts.append(f"""
<table class="layout" style="margin-top:4px;">
<tr>
<td><span style="font-size:11pt;font-weight:bold;color:{primary};">{d.company}</span><span style="font-size:10pt;color:{subtitle_c};">{pos_part}</span></td>
<td style="text-align:right;"><span style="font-size:9pt;color:{date_c};">{date_str}</span></td>
</tr></table>""")
                        if d.description:
                            html_parts.append(f'<div style="font-size:10pt;color:{secondary};margin-top:1px;">{d.description}</div>')
                        if d.achievements:
                            html_parts.append("<ul>" + "".join(f"<li>{a}</li>" for a in d.achievements) + "</ul>")
                    else:
                        html_parts.append(f"""
<table class="layout" style="margin-top:4px;">
<tr>
<td style="width:65%;"><span style="font-size:11pt;font-weight:bold;color:{primary};">{d.company}</span></td>
<td style="width:35%;text-align:right;"><span style="font-size:9pt;color:{date_c};">{date_str}</span></td>
</tr></table>""")
                        if d.position:
                            html_parts.append(f'<div style="font-size:10pt;color:{subtitle_c};">{d.position}</div>')
                        if d.description:
                            html_parts.append(f'<div style="font-size:10pt;color:{secondary};margin-top:2px;">{d.description}</div>')
                        if d.achievements:
                            html_parts.append("<ul>" + "".join(f"<li>{a}</li>" for a in d.achievements) + "</ul>")
                except Exception as e:
                    logger.warning(f"工作经历渲染失败: {e}")

            elif section.type == "project":
                try:
                    d = ProjectItem(**data) if isinstance(data, dict) else data
                    date_str = f"{d.start_date} - {d.end_date or '至今'}" if d.start_date else ""
                    proj_layout = section_layout.get("project", "classic")
                    html_parts.append(f"""
<table class="layout" style="margin-top:4px;">
<tr>
<td style="width:65%;"><span style="font-size:11pt;font-weight:bold;color:{primary};">{d.name}</span></td>
<td style="width:35%;text-align:right;"><span style="font-size:9pt;color:{date_c};">{date_str}</span></td>
</tr></table>""")
                    if d.role:
                        html_parts.append(f'<div style="font-size:10pt;color:{subtitle_c};">{d.role}</div>')
                    if d.tech_stack:
                        if proj_layout == "featured":
                            tech_badges = " ".join(f'<span style="background-color:{accent}18;color:{heading};padding:1px 5px;border-radius:8px;font-size:8pt;margin-right:2px;">{t}</span>' for t in d.tech_stack)
                            html_parts.append(f'<div style="background-color:{accent}10;border:1px solid {accent}25;padding:4px 8px;margin-top:4px;border-radius:4px;"><div style="font-size:8pt;font-weight:bold;color:{heading};text-transform:uppercase;margin-bottom:2px;">Tech Stack</div><div>{tech_badges}</div></div>')
                        else:
                            html_parts.append(f'<div style="font-size:9pt;color:{subtitle_c};margin-top:2px;">技术栈: {" · ".join(d.tech_stack)}</div>')
                    if d.description:
                        html_parts.append(f'<div style="font-size:10pt;color:{secondary};margin-top:2px;">{d.description}</div>')
                    if d.achievements:
                        html_parts.append("<ul>" + "".join(f"<li>{a}</li>" for a in d.achievements) + "</ul>")
                except Exception as e:
                    logger.warning(f"项目经验渲染失败: {e}")

            elif section.type == "skill":
                try:
                    d = SkillItem(**data) if isinstance(data, dict) else data
                    skill_layout = section_layout.get("skill", "inline")
                    if skill_layout == "tags":
                        if d.category:
                            html_parts.append(f'<div style="font-size:10pt;color:{secondary};"><span style="font-weight:bold;color:{subtitle_c};">{d.category} </span>')
                        else:
                            html_parts.append(f'<div style="font-size:10pt;color:{secondary};">')
                        badges = " ".join(f'<span style="background-color:{accent}12;color:{heading};padding:1px 6px;border-radius:10px;font-size:8pt;margin-right:2px;border:1px solid {accent}25;">{s}</span>' for s in d.skills)
                        html_parts.append(badges + '</div>')
                    elif skill_layout == "columns":
                        if d.category:
                            html_parts.append(f'<div style="font-weight:bold;color:{subtitle_c};font-size:10pt;margin-bottom:2px;">{d.category}</div>')
                        cols = 3
                        items = d.skills
                        rows = (len(items) + cols - 1) // cols
                        html_parts.append('<table class="layout" style="margin-top:2px;">')
                        for r in range(rows):
                            html_parts.append('<tr>')
                            for c in range(cols):
                                idx = r * cols + c
                                if idx < len(items):
                                    html_parts.append(f'<td style="width:33%;font-size:10pt;color:{secondary};padding-right:4px;">› {items[idx]}</td>')
                                else:
                                    html_parts.append('<td style="width:33%;"></td>')
                            html_parts.append('</tr>')
                        html_parts.append('</table>')
                    else:
                        html_parts.append(f'<div style="font-size:10pt;color:{secondary};">')
                        if d.category:
                            html_parts.append(f'<span style="font-weight:bold;color:{subtitle_c};">{d.category}: </span>')
                        html_parts.append(f'{" · ".join(d.skills)}</div>')
                except Exception as e:
                    logger.warning(f"技能清单渲染失败: {e}")

            else:
                text = data.get("text", str(data)) if isinstance(data, dict) else str(data)
                html_parts.append(f'<div style="font-size:10pt;color:{secondary};">{text}</div>')

        html_parts.append("</body></html>")
        return "\n".join(html_parts)

    async def generate_pdf(self, resume: ResumeData) -> bytes:
        """通过 docx → PDF 转换生成 PDF（Word 渲染，完美支持中文）"""
        # 先生成 docx
        docx_data = self.generate_docx(resume)

        # 用 docx2pdf 转换为 PDF（利用本地 Word 程序）
        import tempfile
        import subprocess
        tmp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(tmp_dir, "resume.docx")
        pdf_path = os.path.join(tmp_dir, "resume.pdf")
        try:
            with open(docx_path, "wb") as f:
                f.write(docx_data)

            # 使用 docx2pdf 转换
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
            except Exception:
                # fallback: 尝试直接用 Word 命令行转换
                subprocess.run([
                    "powershell",
                    "-Command",
                    f'$w = New-Object -ComObject Word.Application; $w.Visible = $false; '
                    f'$d = $w.Documents.Open("{docx_path}"); '
                    f'$d.SaveAs("{pdf_path}", 17); $d.Close(); $w.Quit()'
                ], check=True, capture_output=True, timeout=30)

            with open(pdf_path, "rb") as f:
                return f.read()
        except Exception as e:
            logger.error(f"docx2pdf 转换失败: {e}")
            # 降级: 返回 html 作为 PDF（使用 xhtml2pdf）
            html_content = self._build_html(resume)
            from xhtml2pdf import pisa
            buffer = io.BytesIO()
            result = pisa.CreatePDF(html_content, dest=buffer, encoding="utf-8")
            if result.err:
                raise RuntimeError(f"PDF 生成失败: {e}")
            buffer.seek(0)
            return buffer.getvalue()
        finally:
            import shutil
            try:
                shutil.rmtree(tmp_dir)
            except Exception:
                pass

    async def export(self, resume: ResumeData, fmt: ExportFormat = ExportFormat.DOCX) -> bytes:
        """导出文档"""
        if fmt == ExportFormat.PDF:
            return await self.generate_pdf(resume)
        return self.generate_docx(resume)


# 全局单例
export_service = ExportService()
